"""
学术摘要提取 API端点 - 全新实现
作者：Ren CBIT https://github.com/reneverland/
"""

import uuid
import io
from typing import Optional
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Body
from fastapi.responses import JSONResponse, StreamingResponse
from loguru import logger
from pydantic import BaseModel

from app.services.academic_extract_service import academic_extract_service
from app.services.document_parser_service import document_parser_service
from app.api.v1.endpoints.auth import get_current_user

router = APIRouter()


# ==================== 数据模型 ====================

class AcademicExtractResponse(BaseModel):
    """学术摘要提取响应"""
    extract_id: str
    summary_zh: str
    summary_en: str
    fact_table: dict
    metadata: dict
    message: str = "提取成功"


class ProgressUpdate(BaseModel):
    """进度更新"""
    stage: str
    progress: int


# ==================== API端点 ====================

@router.post("/extract", response_model=AcademicExtractResponse)
async def extract_academic_article(
    file: UploadFile = File(..., description="学术文章文件（支持DOCX/DOC/PDF）"),
    user: str = Depends(get_current_user)
):
    """
    学术文章智能摘要提取
    
    流程：
    1. 文档获取与预处理 (10%)
    2. 结构化解析与噪音过滤 (20%)
    3. 智能分块与覆盖率保障 (30%)
    4. CBIT-LLM 语义分析 (45%)
    5. 深度理解与事实提取 (60%)
    6. 证据一致性校验 (75%)
    7. 幻觉过滤与可信度评估 (85%)
    8. 摘要生成与风格对齐 (95%)
    9. 准备输送结果 (100%)
    
    返回：
    - summary_zh: 中文摘要（严格4段结构）
    - summary_en: 英文摘要
    - fact_table: 结构化事实表
      - basic_info: 作者、机构、研究问题、数据、方法
      - key_findings: 核心发现清单
      - mechanisms: 作用机制
      - policy_implications: 政策启示
    """
    try:
        filename = file.filename
        logger.info(f"[学术摘要提取] 收到文件: {filename}, 用户: {user}")
        
        # 读取文件内容
        file_content = await file.read()
        
        if not file_content:
            raise HTTPException(status_code=400, detail="文件为空")
        
        # 步骤1: 解析文档
        logger.info("解析文档...")
        try:
            document_data = document_parser_service.parse_document(file_content, filename)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            raise HTTPException(status_code=500, detail=f"文档解析失败: {str(e)}")
        
        # 提取全文文本
        full_text = "\n\n".join([p["text"] for p in document_data["paragraphs"]])
        
        if not full_text.strip():
            raise HTTPException(status_code=400, detail="文档内容为空，无法提取")
        
        # 步骤2: 调用学术摘要提取服务
        logger.info("开始学术摘要提取...")
        
        # 进度回调（在实际应用中可以通过WebSocket推送）
        progress_logs = []
        
        def progress_callback(stage: str, progress: int):
            progress_logs.append({"stage": stage, "progress": progress})
            logger.info(f"[进度更新] {stage}: {progress}%")
        
        try:
            result = await academic_extract_service.extract_academic_article(
                text=full_text,
                progress_callback=progress_callback
            )
        except Exception as e:
            logger.error(f"学术摘要提取失败: {e}")
            import traceback
            logger.error(f"详细错误: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"摘要提取失败: {str(e)}")
        
        # 生成提取ID
        extract_id = str(uuid.uuid4())
        
        logger.info(f"学术摘要提取完成: {filename}")
        logger.info(f"使用模型: {result['metadata'].get('llm_provider')}/{result['metadata'].get('llm_model')}")
        
        return AcademicExtractResponse(
            extract_id=extract_id,
            summary_zh=result["summary_zh"],
            summary_en=result["summary_en"],
            fact_table=result["fact_table"],
            metadata={
                **result["metadata"],
                "filename": filename,
                "user": user,
                "progress_logs": progress_logs
            },
            message="提取成功"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"API处理失败: {e}")
        import traceback
        logger.error(f"详细错误: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"服务器错误: {str(e)}")


@router.get("/progress-stages")
async def get_progress_stages():
    """
    获取进度阶段定义（供前端展示）
    """
    from app.services.academic_extract_service import PROGRESS_STAGES
    return JSONResponse(content={"stages": PROGRESS_STAGES})


@router.get("/health")
async def health_check():
    """
    健康检查
    """
    from app.services.llm_service import llm_service
    
    llm_config = llm_service.get_current_config()
    
    return JSONResponse(content={
        "status": "ok",
        "service": "academic_extract",
        "llm_provider": llm_config.get("provider"),
        "llm_model": llm_config.get("model")
    })


@router.post("/save-to-kb")
async def save_to_knowledge_base(
    data: dict = Body(...),
    user: str = Depends(get_current_user)
):
    """
    保存学术提取结果到知识库（可选）
    
    Body:
    - summary_zh: 中文摘要
    - summary_en: 英文摘要
    - fact_table: 事实表
    """
    try:
        from app.services.knowledge_service import knowledge_service
        import json
        
        logger.info(f"[API Save-to-KB] Received request from user: {user}")
        
        summary_zh = data.get("summary_zh", "")
        summary_en = data.get("summary_en", "")
        fact_table = data.get("fact_table", {})
        tags = data.get("tags", [])
        task_id = data.get("task_id", "")  # 用于关联跳转
        
        logger.info(f"[API Save-to-KB] Data sizes - ZH: {len(summary_zh)} chars, EN: {len(summary_en)} chars")
        
        # 如果没有提供标签，自动推荐
        if not tags:
            try:
                from app.services.knowledge_service import knowledge_service
                content_for_tags = f"{summary_zh}\n{summary_en}"
                tags = await knowledge_service.recommend_tags(content_for_tags, limit=5)
                logger.info(f"[API Save-to-KB] Auto-recommended tags: {tags}")
            except Exception as tag_error:
                logger.warning(f"[API Save-to-KB] Failed to recommend tags: {tag_error}")
                tags = []
        
        # Extract title from fact_table if available
        title = "Academic Extract"
        if fact_table.get("basic_info", {}).get("research_question"):
            title = fact_table["basic_info"]["research_question"].get("question", "Academic Extract")[:100]
        
        logger.info(f"[API Save-to-KB] Title: {title}")
        
        kb_metadata = {
            "title": title,
            "type": "academic_extract",
            "summary_zh": summary_zh,
            "summary_en": summary_en,
            "fact_table": fact_table,
            "tags": tags,
            "source_task_id": task_id,  # 用于关联跳转
            "source_type": "academic_extract",  # 用于识别来源类型
            "user": user
        }
        
        # Combine content for indexing
        tags_str = ", ".join(tags)
        kb_content = f"Title: {title}\nTags: {tags_str}\nSummary (ZH): {summary_zh}\nSummary (EN): {summary_en}\nFindings: {json.dumps(fact_table.get('key_findings', []), ensure_ascii=False)}"
        
        logger.info(f"[API Save-to-KB] Content prepared. Total length: {len(kb_content)} chars. Calling knowledge_service...")
        
        await knowledge_service.add_document(content=kb_content, metadata=kb_metadata)
        
        logger.info(f"[API Save-to-KB] ✅ Successfully saved to Knowledge Base: {title} (User: {user})")
        
        return JSONResponse(content={
            "status": "success",
            "message": "已成功保存到知识库",
            "title": title
        })
        
    except Exception as e:
        logger.error(f"[API Save-to-KB] ❌ Failed to save to Knowledge Base: {e}")
        import traceback
        logger.error(f"[API Save-to-KB] Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"保存失败: {str(e)}")


@router.get("/extracts")
async def get_academic_extracts(
    limit: int = 20,
    offset: int = 0,
    search: str = None,
    user: str = Depends(get_current_user)
):
    """
    获取学术提取历史记录
    
    从知识库中获取该用户的 academic extract 记录
    支持搜索和分页
    """
    try:
        from app.services.knowledge_service import knowledge_service
        
        # 检查知识库服务是否可用
        if not knowledge_service.initialized:
            logger.warning(f"[API Get-Extracts] KnowledgeService not initialized, returning empty list")
            return JSONResponse(content={
                "items": [],
                "total": 0,
                "limit": limit,
                "offset": offset,
                "warning": "知识库服务暂时不可用"
            })
        
        logger.info(f"[API Get-Extracts] Fetching academic extracts for user: {user}, limit: {limit}, offset: {offset}, search: {search}")
        
        # 从知识库获取文档（获取更多以便过滤和搜索）
        docs = await knowledge_service.list_documents(limit=100, offset=0)
        
        # 过滤出 academic_extract 类型的文档
        academic_extracts = []
        for doc in docs:
            metadata = doc.get("payload", {})
            if metadata.get("type") == "academic_extract":
                # 如果有搜索词，进行过滤
                if search:
                    search_lower = search.lower()
                    title = metadata.get("title", "").lower()
                    summary_zh = metadata.get("summary_zh", "").lower()
                    summary_en = metadata.get("summary_en", "").lower()
                    
                    if search_lower not in title and search_lower not in summary_zh and search_lower not in summary_en:
                        continue
                
                # 提取关键信息
                extract_item = {
                    "id": doc.get("id"),
                    "title": metadata.get("title", "Untitled"),
                    "summary_zh": metadata.get("summary_zh", "")[:200],  # 截取前200字符作为预览
                    "summary_en": metadata.get("summary_en", "")[:200],
                    "created_at": metadata.get("created_at", ""),
                    "user": metadata.get("user", ""),
                    "metadata": metadata
                }
                academic_extracts.append(extract_item)
        
        # 应用分页
        total = len(academic_extracts)
        paginated_extracts = academic_extracts[offset:offset+limit]
        
        logger.info(f"[API Get-Extracts] Found {total} academic extracts, returning {len(paginated_extracts)}")
        
        return JSONResponse(content={
            "items": paginated_extracts,
            "total": total,
            "limit": limit,
            "offset": offset
        })
        
    except Exception as e:
        logger.error(f"[API Get-Extracts] Failed to fetch extracts: {e}")
        import traceback
        logger.error(f"[API Get-Extracts] Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"获取历史记录失败: {str(e)}")


@router.get("/extracts/{extract_id}")
async def get_academic_extract_detail(
    extract_id: str,
    user: str = Depends(get_current_user)
):
    """
    获取学术提取详情
    
    根据 ID 获取完整的提取结果
    """
    try:
        from app.services.knowledge_service import knowledge_service
        
        # 检查知识库服务是否可用
        if not knowledge_service.initialized:
            logger.warning(f"[API Get-Extract-Detail] KnowledgeService not initialized")
            raise HTTPException(status_code=503, detail="知识库服务暂时不可用")
        
        logger.info(f"[API Get-Extract-Detail] Fetching extract {extract_id} for user: {user}")
        
        # 从知识库获取文档详情
        doc = await knowledge_service.get_document(extract_id)
        
        if not doc:
            raise HTTPException(status_code=404, detail="记录不存在")
        
        metadata = doc.get("payload", {})
        
        # 验证权限（确保是该用户的记录）
        if metadata.get("user") != user and user != "admin":
            raise HTTPException(status_code=403, detail="无权访问此记录")
        
        logger.info(f"[API Get-Extract-Detail] Successfully fetched extract {extract_id}")
        
        return JSONResponse(content={
            "id": doc.get("id"),
            "title": metadata.get("title", "Untitled"),
            "summary_zh": metadata.get("summary_zh", ""),
            "summary_en": metadata.get("summary_en", ""),
            "fact_table": metadata.get("fact_table", {}),
            "created_at": metadata.get("created_at", ""),
            "user": metadata.get("user", ""),
            "metadata": metadata
        })
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[API Get-Extract-Detail] Failed to fetch detail: {e}")
        import traceback
        logger.error(f"[API Get-Extract-Detail] Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"获取详情失败: {str(e)}")


@router.delete("/extracts/{extract_id}")
async def delete_academic_extract(
    extract_id: str,
    user: str = Depends(get_current_user)
):
    """
    删除学术提取记录
    """
    try:
        from app.services.knowledge_service import knowledge_service
        
        # 检查知识库服务是否可用
        if not knowledge_service.initialized:
            logger.warning(f"[API Delete-Extract] KnowledgeService not initialized")
            raise HTTPException(status_code=503, detail="知识库服务暂时不可用")
        
        logger.info(f"[API Delete-Extract] Deleting extract {extract_id} for user: {user}")
        
        # 先获取文档验证权限
        doc = await knowledge_service.get_document(extract_id)
        
        if not doc:
            raise HTTPException(status_code=404, detail="记录不存在")
        
        metadata = doc.get("payload", {})
        
        # 验证权限
        if metadata.get("user") != user and user != "admin":
            raise HTTPException(status_code=403, detail="无权删除此记录")
        
        # 删除文档
        success = await knowledge_service.delete_document(extract_id)
        
        if success:
            logger.info(f"[API Delete-Extract] Successfully deleted extract {extract_id}")
            return JSONResponse(content={"status": "success", "message": "删除成功"})
        else:
            raise HTTPException(status_code=500, detail="删除失败")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[API Delete-Extract] Failed to delete: {e}")
        import traceback
        logger.error(f"[API Delete-Extract] Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"删除失败: {str(e)}")


@router.post("/recommend-tags")
async def recommend_tags_for_extract(
    data: dict = Body(...),
    user: str = Depends(get_current_user)
):
    """
    为学术摘要推荐标签
    
    Body:
    - text: 摘要内容
    - limit: 返回标签数量（默认5）
    """
    try:
        from app.services.knowledge_service import knowledge_service
        
        text = data.get("text", "")
        limit = data.get("limit", 5)
        
        if not text:
            raise HTTPException(status_code=400, detail="缺少 text 参数")
        
        tags = await knowledge_service.recommend_tags(text[:2000], limit=limit)
        
        return JSONResponse(content={
            "tags": tags,
            "count": len(tags)
        })
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Academic Recommend-Tags] Failed: {e}")
        raise HTTPException(status_code=500, detail=f"推荐标签失败: {str(e)}")


@router.post("/download-word")
async def download_as_word(
    data: dict = Body(...),
    user: str = Depends(get_current_user)
):
    """
    将学术摘要下载为 Word 文档
    
    Body:
    - summary_zh: 中文摘要
    - summary_en: 英文摘要
    - fact_table: 事实表
    - title: 文档标题（可选）
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import json
        
        summary_zh = data.get("summary_zh", "")
        summary_en = data.get("summary_en", "")
        fact_table = data.get("fact_table", {})
        title = data.get("title", "学术摘要提取结果")
        
        if not summary_zh and not summary_en:
            raise HTTPException(status_code=400, detail="缺少摘要内容")
        
        # 创建 Word 文档
        doc = Document()
        
        # 设置标题
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 中文摘要
        if summary_zh:
            doc.add_heading("中文摘要", level=1)
            for para in summary_zh.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # 英文摘要
        if summary_en:
            doc.add_heading("English Summary", level=1)
            for para in summary_en.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # 事实表
        if fact_table:
            doc.add_heading("结构化事实表", level=1)
            
            # 基本信息
            basic_info = fact_table.get("basic_info", {})
            if basic_info:
                doc.add_heading("基本信息", level=2)
                for key, value in basic_info.items():
                    if isinstance(value, dict):
                        doc.add_paragraph(f"{key}: {json.dumps(value, ensure_ascii=False)}")
                    elif isinstance(value, list):
                        doc.add_paragraph(f"{key}: {', '.join(str(v) for v in value)}")
                    else:
                        doc.add_paragraph(f"{key}: {value}")
            
            # 核心发现
            key_findings = fact_table.get("key_findings", [])
            if key_findings:
                doc.add_heading("核心发现", level=2)
                for i, finding in enumerate(key_findings, 1):
                    if isinstance(finding, dict):
                        doc.add_paragraph(f"{i}. {finding.get('finding', finding)}")
                    else:
                        doc.add_paragraph(f"{i}. {finding}")
        
        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # 返回文件 - 使用 URL 编码处理中文文件名
        from urllib.parse import quote
        safe_title = title.replace('/', '_').replace('\\', '_')[:50]  # 限制长度并移除非法字符
        filename_encoded = quote(f"{safe_title}.docx")
        
        logger.info(f"[Download Word] Generating document: {safe_title}.docx")
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=\"document.docx\"; filename*=UTF-8''{filename_encoded}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Download Word] Failed: {e}")
        import traceback
        logger.error(f"[Download Word] Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"生成Word文档失败: {str(e)}")
