"""
Integrated Voiceover API Endpoints
多文献整合口播稿件策划接口
"""
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import StreamingResponse
from typing import List, Optional
from loguru import logger
import io
from app.models.schemas import (
    IntegratedVoiceoverRequest,
    IntegratedVoiceoverResponse,
    IntegratedVoiceoverStatus
)
from app.services.integrated_voiceover_service import integrated_voiceover_service
from app.api.v1.endpoints.auth import get_current_user

router = APIRouter()


@router.post("/create", response_model=dict)
async def create_integrated_voiceover_task(
    topic_hint: str = Form(..., description="主题/问题一句话"),
    speaker_affiliation: Optional[str] = Form(None, description="主播机构"),
    speaker_name: Optional[str] = Form(None, description="主播姓名"),
    include_vox_intro: bool = Form(True, description="是否包含VOXCHINA片头"),
    style_preference: Optional[str] = Form(None, description="结构偏好: S1/S2/S3/S4"),
    language: str = Form("zh", description="语言"),
    word_limit: Optional[int] = Form(None, description="字数限制: 2000/3000/4000 或不填(不限制)"),
    files: List[UploadFile] = File(..., description="上传的文档（Word/PDF）"),
    current_user: dict = Depends(get_current_user)
):
    """
    创建整合口播稿生成任务
    
    上传多个文档，系统将：
    1. 解析文档内容（文字、表格、图表）
    2. 提取证据和图表
    3. 生成结构化口播稿
    
    返回task_id，可用于查询任务状态和结果
    """
    try:
        # 验证文件
        if not files or len(files) == 0:
            raise HTTPException(status_code=400, detail="至少需要上传一个文档")
        
        # 验证文件格式
        allowed_extensions = ['.docx', '.doc', '.pdf']
        for file in files:
            file_ext = '.' + file.filename.split('.')[-1].lower()
            if file_ext not in allowed_extensions:
                raise HTTPException(
                    status_code=400,
                    detail=f"不支持的文件格式: {file.filename}，仅支持 .docx, .doc, .pdf"
                )
        
        # 读取文件内容
        file_data = []
        for file in files:
            content = await file.read()
            file_data.append((file.filename, content))
        
        # 创建请求对象
        request = IntegratedVoiceoverRequest(
            topic_hint=topic_hint,
            speaker_affiliation=speaker_affiliation,
            speaker_name=speaker_name,
            include_vox_intro=include_vox_intro,
            style_preference=style_preference,
            language=language,
            word_limit=word_limit
        )
        
        # 创建任务
        task_id = await integrated_voiceover_service.create_task(request, file_data)
        
        logger.info(f"Created integrated voiceover task: {task_id} by user {current_user.get('username')}")
        
        return {
            "task_id": task_id,
            "message": "任务创建成功，正在处理中...",
            "status": "processing"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to create integrated voiceover task: {e}")
        raise HTTPException(status_code=500, detail=f"创建任务失败: {str(e)}")


@router.get("/status/{task_id}")
async def get_task_status(
    task_id: str,
    current_user: dict = Depends(get_current_user)
):
    """
    查询任务状态
    
    返回任务的当前状态、进度和当前步骤
    包含 parsed_docs 用于文档详情查看
    """
    try:
        logger.info(f"Getting task status for: {task_id}")
        task_data = integrated_voiceover_service.get_task_status(task_id)
        
        if not task_data:
            logger.warning(f"Task not found: {task_id}")
            raise HTTPException(status_code=404, detail="任务不存在")
        
        logger.info(f"Task {task_id} status: {task_data['status']}")
        
        # 如果任务完成，返回完整结果
        result = None
        if task_data["status"] == "completed":
            logger.info(f"Task {task_id} is completed, fetching result...")
            try:
                result_obj = integrated_voiceover_service.get_task_result(task_id)
                if result_obj:
                    # 转换 Pydantic 模型为字典
                    result = result_obj.dict() if hasattr(result_obj, 'dict') else result_obj.model_dump()
                logger.info(f"✅ Task result fetched successfully for {task_id}")
            except Exception as result_error:
                logger.error(f"❌ Error fetching task result for {task_id}: {result_error}", exc_info=True)
                raise HTTPException(status_code=500, detail=f"获取任务结果失败: {str(result_error)}")
        
        # 构建响应，包含 parsed_docs
        response = {
            "task_id": task_id,
            "status": task_data["status"],
            "progress": task_data["progress"],
            "current_step": task_data["current_step"],
            "error": task_data.get("error"),
            "result": result,  # 已经是字典格式
            "parsed_docs": task_data.get("parsed_docs", [])  # 添加 parsed_docs
        }
        
        logger.info(f"✅ Returning response for {task_id}, result present: {result is not None}")
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get task status for {task_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"查询任务状态失败: {str(e)}")


@router.get("/result/{task_id}", response_model=IntegratedVoiceoverResponse)
async def get_task_result(
    task_id: str,
    current_user: dict = Depends(get_current_user)
):
    """
    获取任务结果
    
    返回完整的口播稿生成结果，包括：
    - 风格配置
    - 证据台账
    - 图表台账
    - 结构设计
    - 审阅版口播稿
    - 上屏版口播稿
    """
    try:
        task_data = integrated_voiceover_service.get_task_status(task_id)
        
        if not task_data:
            raise HTTPException(status_code=404, detail="任务不存在")
        
        if task_data["status"] == "processing":
            raise HTTPException(status_code=400, detail="任务还在处理中，请稍后再试")
        
        if task_data["status"] == "failed":
            error_msg = task_data.get("error", "未知错误")
            raise HTTPException(status_code=500, detail=f"任务处理失败: {error_msg}")
        
        result = integrated_voiceover_service.get_task_result(task_id)
        
        if not result:
            raise HTTPException(status_code=404, detail="任务结果不存在")
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get task result: {e}")
        raise HTTPException(status_code=500, detail=f"获取任务结果失败: {str(e)}")


@router.get("/list", response_model=List[dict])
async def list_tasks(
    current_user: dict = Depends(get_current_user)
):
    """
    列出所有任务
    
    返回用户的所有任务列表（简化信息）
    """
    try:
        from datetime import datetime
        
        # 获取所有任务
        tasks = []
        for task_id, task_data in integrated_voiceover_service.tasks.items():
            # 统一转换日期格式为 ISO 字符串
            created_at = task_data["created_at"]
            if isinstance(created_at, datetime):
                created_at = created_at.isoformat()
            
            updated_at = task_data["updated_at"]
            if isinstance(updated_at, datetime):
                updated_at = updated_at.isoformat()
            
            tasks.append({
                "task_id": task_id,
                "status": task_data["status"],
                "progress": task_data["progress"],
                "current_step": task_data["current_step"],
                "topic_hint": task_data["request"].get("topic_hint", ""),
                "files_count": len(task_data.get("files", [])),
                "created_at": created_at,
                "updated_at": updated_at
            })
        
        # 按创建时间倒序排序
        tasks.sort(key=lambda x: x["created_at"], reverse=True)
        
        logger.info(f"✅ 成功获取 {len(tasks)} 个任务列表")
        return tasks
        
    except Exception as e:
        logger.error(f"Failed to list tasks: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"获取任务列表失败: {str(e)}")


@router.delete("/delete/{task_id}")
async def delete_task(
    task_id: str,
    current_user: dict = Depends(get_current_user)
):
    """
    删除任务
    
    删除指定的任务及其所有数据
    """
    try:
        if task_id not in integrated_voiceover_service.tasks:
            raise HTTPException(status_code=404, detail="任务不存在")
        
        del integrated_voiceover_service.tasks[task_id]
        
        # 保存任务状态
        integrated_voiceover_service._save_tasks()
        
        logger.info(f"Deleted task {task_id} by user {current_user.get('username')}")
        
        return {
            "message": "任务已删除",
            "task_id": task_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to delete task: {e}")
        raise HTTPException(status_code=500, detail=f"删除任务失败: {str(e)}")


@router.post("/save-to-kb")
async def save_to_knowledge_base(
    data: dict,
    current_user: dict = Depends(get_current_user)
):
    """
    保存整合口播稿结果到知识库
    
    Body:
    - script_final: 最终脚本
    - topic_hint: 主题提示
    - speaker_name: 主播姓名
    - speaker_affiliation: 主播机构
    - tags: 标签列表（可选，如果不提供则自动推荐）
    - task_id: 任务ID（可选，用于关联跳转）
    """
    try:
        from app.services.knowledge_service import knowledge_service
        
        logger.info(f"[Integrated Save-to-KB] Received request from user: {current_user.get('username')}")
        
        script_final = data.get("script_final", "")
        topic_hint = data.get("topic_hint", "")
        speaker_name = data.get("speaker_name", "")
        speaker_affiliation = data.get("speaker_affiliation", "")
        tags = data.get("tags", [])
        task_id = data.get("task_id", "")
        
        if not script_final:
            raise HTTPException(status_code=400, detail="缺少 script_final 参数")
        
        logger.info(f"[Integrated Save-to-KB] Script length: {len(script_final)} chars")
        
        # 准备标题
        title = topic_hint if topic_hint else "整合口播稿"
        if speaker_name:
            title = f"{title} - {speaker_name}"
        
        logger.info(f"[Integrated Save-to-KB] Title: {title}")
        
        # 如果没有提供标签，自动推荐
        if not tags:
            try:
                tags = await knowledge_service.recommend_tags(script_final[:2000], limit=5)
                logger.info(f"[Integrated Save-to-KB] Auto-recommended tags: {tags}")
            except Exception as tag_error:
                logger.warning(f"[Integrated Save-to-KB] Failed to recommend tags: {tag_error}")
                tags = []
        
        kb_metadata = {
            "title": title,
            "type": "integrated_voiceover",
            "topic_hint": topic_hint,
            "speaker_name": speaker_name,
            "speaker_affiliation": speaker_affiliation,
            "tags": tags,
            "source_task_id": task_id,  # 用于关联跳转
            "source_type": "integrated_voiceover",  # 用于识别来源类型
            "user": current_user.get("username")
        }
        
        # 保存到知识库
        logger.info(f"[Integrated Save-to-KB] Calling knowledge_service...")
        
        doc_id = await knowledge_service.add_document(
            content=script_final, 
            metadata=kb_metadata
        )
        
        logger.info(f"[Integrated Save-to-KB] ✅ Successfully saved to Knowledge Base: {title} (Doc ID: {doc_id})")
        
        return {
            "status": "success",
            "message": "已成功保存到知识库",
            "title": title,
            "doc_id": doc_id,
            "tags": tags
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Integrated Save-to-KB] ❌ Failed to save to Knowledge Base: {e}")
        import traceback
        logger.error(f"[Integrated Save-to-KB] Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"保存失败: {str(e)}")


@router.post("/recommend-tags")
async def recommend_tags_for_voiceover(
    data: dict,
    current_user: dict = Depends(get_current_user)
):
    """
    为口播稿推荐标签
    
    Body:
    - text: 口播稿内容
    - limit: 返回标签数量（默认5）
    """
    try:
        from app.services.knowledge_service import knowledge_service
        
        text = data.get("text", "")
        limit = data.get("limit", 5)
        
        if not text:
            raise HTTPException(status_code=400, detail="缺少 text 参数")
        
        tags = await knowledge_service.recommend_tags(text[:2000], limit=limit)
        
        return {
            "tags": tags,
            "count": len(tags)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Integrated Recommend-Tags] Failed: {e}")
        raise HTTPException(status_code=500, detail=f"推荐标签失败: {str(e)}")


@router.post("/download-word")
async def download_as_word(
    data: dict,
    current_user: dict = Depends(get_current_user)
):
    """
    将口播稿下载为 Word 文档
    
    Body:
    - content: 口播稿内容
    - title: 文档标题（可选）
    - image_assets: 图片资产列表（可选）
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import os
        import re
        
        content = data.get("content", "")
        title = data.get("title", "VoxChina口播稿")
        image_assets = data.get("image_assets", [])
        
        if not content:
            raise HTTPException(status_code=400, detail="缺少 content 参数")
        
        # 构建图片资产映射 {asset_id: image_path}
        image_map = {}
        for asset in image_assets:
            if asset.get("image_path") and asset.get("asset_id"):
                image_map[asset["asset_id"]] = {
                    "path": asset["image_path"],
                    "caption": asset.get("caption", asset["asset_id"])
                }
        
        logger.info(f"[Download Word] Image assets: {len(image_map)}")
        
        # 创建 Word 文档
        doc = Document()
        
        # 设置标题
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 用于追踪已插入的图片
        inserted_images = set()
        
        # 解析内容并添加到文档
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 处理标题
            if line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            else:
                # 普通段落
                para = doc.add_paragraph(line)
                para.style.font.size = Pt(12)
                
                # 检查是否有图片引用 (D1-FIG-1, D2-FIG-2 等格式)
                fig_refs = re.findall(r'D\d+-FIG-\d+', line)
                for fig_ref in fig_refs:
                    if fig_ref in image_map and fig_ref not in inserted_images:
                        img_info = image_map[fig_ref]
                        img_path = img_info["path"]
                        
                        # 检查文件是否存在
                        if os.path.exists(img_path):
                            try:
                                # 添加图片
                                doc.add_picture(img_path, width=Inches(5.5))
                                # 添加图片说明
                                caption_para = doc.add_paragraph(f"图: {img_info['caption']}")
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption_para.style.font.size = Pt(10)
                                inserted_images.add(fig_ref)
                                logger.info(f"[Download Word] Inserted image: {fig_ref}")
                            except Exception as img_error:
                                logger.warning(f"[Download Word] Failed to insert image {fig_ref}: {img_error}")
                        else:
                            logger.warning(f"[Download Word] Image file not found: {img_path}")
        
        # 如果有未插入的图片，在文档末尾添加附录
        remaining_images = set(image_map.keys()) - inserted_images
        if remaining_images:
            doc.add_heading("附录：图表", level=1)
            for fig_ref in sorted(remaining_images):
                img_info = image_map[fig_ref]
                img_path = img_info["path"]
                if os.path.exists(img_path):
                    try:
                        doc.add_heading(img_info['caption'], level=2)
                        doc.add_picture(img_path, width=Inches(5.5))
                        logger.info(f"[Download Word] Inserted appendix image: {fig_ref}")
                    except Exception as img_error:
                        logger.warning(f"[Download Word] Failed to insert appendix image {fig_ref}: {img_error}")
        
        logger.info(f"[Download Word] Total images inserted: {len(inserted_images)}")
        
        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # 返回文件 - 使用 URL 编码处理中文文件名
        from urllib.parse import quote
        safe_title = title.replace('/', '_').replace('\\', '_')[:50]  # 限制长度并移除非法字符
        filename_encoded = quote(f"{safe_title}.docx")
        
        logger.info(f"[Download Word] Generating document: {safe_title}.docx")
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=\"document.docx\"; filename*=UTF-8''{filename_encoded}"
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[Download Word] Failed: {e}")
        import traceback
        logger.error(f"[Download Word] Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"生成Word文档失败: {str(e)}")
