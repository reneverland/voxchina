"""
口播稿生成 - DOCX导出服务
将最终稿件导出为格式化的Word文档
"""
import os
import re
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger


class ScriptDocxService:
    """DOCX导出服务"""
    
    def __init__(self):
        self.output_dir = "backend/static/scripts"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def _add_title(self, doc: Document, title: str):
        """添加标题段落"""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.name = "Microsoft YaHei"
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(12)
    
    def _add_paragraph(self, doc: Document, text: str, style: str = "normal"):
        """添加普通段落"""
        para = doc.add_paragraph(text)
        
        if style == "heading2":
            run = para.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = "Microsoft YaHei"
            para.space_before = Pt(12)
            para.space_after = Pt(6)
        elif style == "figure":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True
            para.space_before = Pt(6)
            para.space_after = Pt(6)
        else:
            # 普通段落
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Microsoft YaHei"
            para.space_after = Pt(8)
            para.line_spacing = 1.5
    
    def _parse_script_structure(self, script_text: str) -> Dict[str, Any]:
        """
        解析稿件结构
        
        Returns:
            {
                "title": str,
                "sections": [
                    {"type": "title|paragraph|heading2|figure", "text": str}
                ]
            }
        """
        lines = script_text.split("\n")
        sections = []
        title = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 第一个非空行作为标题
            if title is None:
                title = line
                sections.append({"type": "title", "text": line})
                continue
            
            # 检测小标题（通常以数字或"一、二、三"开头，或全大写/加粗标记）
            if re.match(r'^(一|二|三|四|五|六|七|八|九|十|[\d]+[、\.])\s*.+', line):
                sections.append({"type": "heading2", "text": line})
            # 检测Figure占位行
            elif line.startswith("Figure") or line.startswith("图"):
                sections.append({"type": "figure", "text": line})
            else:
                sections.append({"type": "paragraph", "text": line})
        
        return {
            "title": title or "VoxChina口播稿",
            "sections": sections
        }
    
    def export_to_docx(
        self,
        script_text: str,
        task_id: str,
        episode_title: str = None
    ) -> str:
        """
        导出为DOCX文件
        
        Args:
            script_text: 完整口播稿文本
            task_id: 任务ID
            episode_title: 期数标题
            
        Returns:
            文件路径（相对路径，用于下载）
        """
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置默认字体
            style = doc.styles['Normal']
            style.font.name = "Microsoft YaHei"
            style.font.size = Pt(11)
            
            # 解析稿件结构
            structure = self._parse_script_structure(script_text)
            
            # 添加内容
            for section in structure["sections"]:
                section_type = section["type"]
                section_text = section["text"]
                
                if section_type == "title":
                    self._add_title(doc, section_text)
                elif section_type == "heading2":
                    self._add_paragraph(doc, section_text, style="heading2")
                elif section_type == "figure":
                    self._add_paragraph(doc, section_text, style="figure")
                else:
                    self._add_paragraph(doc, section_text, style="normal")
            
            # 生成文件名
            date_str = datetime.now().strftime("%Y%m%d")
            safe_title = (episode_title or structure["title"])[:20]
            safe_title = re.sub(r'[^\w\s-]', '', safe_title).strip()
            filename = f"AI专题口播稿件_{safe_title}_{date_str}_{task_id[:8]}.docx"
            
            # 保存文件
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            
            logger.info(f"Exported DOCX: {filepath}")
            
            # 返回相对路径（用于前端下载）
            return f"/static/scripts/{filename}"
            
        except Exception as e:
            logger.error(f"Failed to export DOCX: {e}")
            raise Exception(f"DOCX导出失败: {str(e)}")


script_docx_service = ScriptDocxService()

