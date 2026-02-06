"""
æ–‡æ¡£è§£ææœåŠ¡ - æ”¯æŒ DOCX/DOC/PDF
å°†æ–‡æ¡£è½¬æ¢ä¸ºç»“æ„åŒ–JSONï¼Œä¿ç•™æ®µè½åºå·å’Œå±‚æ¬¡ç»“æ„
"""
import re
from typing import List, Dict, Any
from loguru import logger
from docx import Document
from pypdf import PdfReader
import io
import uuid
from pathlib import Path


class DocumentParserService:
    """æ–‡æ¡£è§£ææœåŠ¡"""
    
    def __init__(self):
        # è¿‡æ»¤è§„åˆ™ï¼šæ˜æ˜¾æ— ä¿¡æ¯çš„æ®µè½æ¨¡å¼
        self.noise_patterns = [
            r'^copyright\s*Â©',
            r'^Â©\s*\d{4}',
            r'^email\s*:',
            r'^e-mail\s*:',
            r'^tel\s*:',
            r'^phone\s*:',
            r'^fax\s*:',
            r'^\s*$',  # ç©ºè¡Œ
            r'^page\s+\d+',
            r'^\d+\s*/\s*\d+$',  # é¡µç 
        ]
        
    def is_noise(self, text: str) -> bool:
        """åˆ¤æ–­æ˜¯å¦ä¸ºå™ªéŸ³æ®µè½"""
        text_lower = text.strip().lower()
        
        # ç©ºæ®µè½
        if not text_lower:
            return True
        
        # å¤ªçŸ­çš„æ®µè½ï¼ˆå°äº5ä¸ªå­—ç¬¦ï¼‰
        if len(text_lower) < 5:
            return True
        
        # åŒ¹é…å™ªéŸ³æ¨¡å¼
        for pattern in self.noise_patterns:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True
        
        return False
    
    def extract_images_from_docx(self, doc_obj) -> List[Dict[str, str]]:
        """ä»æ–‡æ¡£å¯¹è±¡ä¸­æå–å›¾ç‰‡å¹¶ä¿å­˜"""
        images = []
        try:
            # è·å–å½“å‰æ–‡ä»¶æ‰€åœ¨ç›®å½•ï¼Œå‘ä¸Šä¸¤çº§åˆ°è¾¾backendç›®å½•
            current_file = Path(__file__)
            backend_dir = current_file.parent.parent.parent  # backend/
            save_dir = backend_dir / "static" / "images" / "extracted"
            
            # ç¡®ä¿ç›®å½•å­˜åœ¨
            try:
                save_dir.mkdir(parents=True, exist_ok=True)
                logger.info(f"å›¾ç‰‡ä¿å­˜ç›®å½•: {save_dir}")
            except Exception as e:
                logger.error(f"æ— æ³•åˆ›å»ºå›¾ç‰‡ä¿å­˜ç›®å½•: {e}")
                # å°è¯•å¤‡ç”¨è·¯å¾„
                try:
                    save_dir = Path("/www/wwwroot/voxchina/backend/static/images/extracted")
                    save_dir.mkdir(parents=True, exist_ok=True)
                    logger.info(f"ä½¿ç”¨å¤‡ç”¨å›¾ç‰‡ä¿å­˜ç›®å½•: {save_dir}")
                except Exception as e2:
                    logger.error(f"æ— æ³•åˆ›å»ºå¤‡ç”¨å›¾ç‰‡ä¿å­˜ç›®å½•: {e2}")
                    return []
            
            # éå†æ‰€æœ‰å…³ç³»å¯»æ‰¾å›¾ç‰‡
            processed_blobs = set()
            
            for rel in doc_obj.part.rels.values():
                if "image" in rel.target_ref:
                    if hasattr(rel, "target_part"):
                        try:
                            image_data = rel.target_part.blob
                            # ä½¿ç”¨å†…å®¹å“ˆå¸Œå»é‡
                            import hashlib
                            img_hash = hashlib.md5(image_data).hexdigest()
                            if img_hash in processed_blobs:
                                continue
                            processed_blobs.add(img_hash)

                            content_type = rel.target_part.content_type
                            
                            # ç¡®å®šæ‰©å±•å
                            ext = ".jpg"
                            if "png" in content_type: ext = ".png"
                            elif "jpeg" in content_type: ext = ".jpg"
                            elif "gif" in content_type: ext = ".gif"
                            
                            # ç”Ÿæˆæ–‡ä»¶å
                            img_filename = f"img_{uuid.uuid4().hex[:12]}{ext}"
                            img_path = save_dir / img_filename
                            
                            # ä¿å­˜æ–‡ä»¶
                            with open(img_path, "wb") as f:
                                f.write(image_data)
                                
                            images.append({
                                "url": f"/static/images/extracted/{img_filename}",
                                "path": str(img_path),
                                "filename": img_filename
                            })
                        except Exception as img_err:
                            logger.warning(f"Failed to save extracted image: {img_err}")
                            continue
                            
            logger.info(f"Extracted {len(images)} images from document")
        except Exception as e:
            logger.error(f"æå–å›¾ç‰‡æ•´ä½“æµç¨‹å¤±è´¥: {e}")
            
        return images

    def parse_docx(self, file_content: bytes) -> Dict[str, Any]:
        """è§£æDOCXæ–‡æ¡£"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            # æå–å›¾ç‰‡
            extracted_images = self.extract_images_from_docx(doc)
            
            title = ""
            paragraphs = []
            paragraph_id = 0
            
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                # è·³è¿‡å™ªéŸ³æ®µè½
                if self.is_noise(text):
                    continue
                
                # ç¬¬ä¸€ä¸ªéç©ºæ®µè½ä½œä¸ºæ ‡é¢˜
                if not title and text:
                    title = text
                    continue
                
                # æå–æ®µè½æ ·å¼ä¿¡æ¯ï¼ˆæ ‡é¢˜çº§åˆ«ï¼‰
                style_name = para.style.name if para.style else "Normal"
                is_heading = "Heading" in style_name
                heading_level = 0
                
                if is_heading:
                    # æå–æ ‡é¢˜çº§åˆ« (Heading 1 -> level 1)
                    match = re.search(r'Heading\s+(\d+)', style_name)
                    if match:
                        heading_level = int(match.group(1))
                
                paragraphs.append({
                    "paragraph_id": paragraph_id,
                    "text": text,
                    "type": "heading" if is_heading else "paragraph",
                    "heading_level": heading_level,
                    "original_index": idx
                })
                
                paragraph_id += 1
            
            # å¤„ç†è¡¨æ ¼å†…å®¹
            table_count = 0
            for table_idx, table in enumerate(doc.tables):
                table_count += 1
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text and not self.is_noise(row_text):
                        paragraphs.append({
                            "paragraph_id": paragraph_id,
                            "text": row_text,
                            "type": "table_row",
                            "heading_level": 0,
                            "original_index": f"table_{table_idx}"
                        })
                        paragraph_id += 1
            
            logger.info(f"âœ… Parsed DOCX: {len(paragraphs)} paragraphs, {table_count} tables, {len(extracted_images)} images, title: {title[:50]}")
            
            # è¯¦ç»†è®°å½•å›¾ç‰‡ä¿¡æ¯
            if extracted_images:
                logger.info(f"ğŸ“¸ æå–åˆ° {len(extracted_images)} å¼ å›¾ç‰‡:")
                for idx, img in enumerate(extracted_images):
                    logger.info(f"   å›¾ç‰‡ {idx+1}: {img.get('filename')} - URL: {img.get('url')}")
            else:
                logger.warning("âš ï¸  æœªæå–åˆ°ä»»ä½•å›¾ç‰‡")
            
            result = {
                "title": title,
                "total_paragraphs": len(paragraphs),
                "total_tables": table_count,  # è¡¨æ ¼æ•°é‡
                "paragraphs": paragraphs,
                "images": extracted_images,  # æ·»åŠ æå–çš„å›¾ç‰‡
                "format": "docx"
            }
            
            logger.info(f"ğŸ“¦ è¿”å›çš„ result åŒ…å« images å­—æ®µ: {'images' in result}, å€¼: {len(result.get('images', []))} å¼ å›¾ç‰‡")
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise Exception(f"DOCXè§£æå¤±è´¥: {str(e)}")
    
    def parse_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """è§£æPDFæ–‡æ¡£"""
        try:
            pdf_reader = PdfReader(io.BytesIO(file_content))
            
            title = ""
            paragraphs = []
            paragraph_id = 0
            
            # å°è¯•ä»metadataè·å–æ ‡é¢˜
            if pdf_reader.metadata and pdf_reader.metadata.title:
                title = pdf_reader.metadata.title
            
            all_text = []
            
            # é€é¡µæå–æ–‡æœ¬
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        all_text.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue
            
            # åˆå¹¶æ‰€æœ‰é¡µé¢æ–‡æœ¬
            full_text = "\n".join(all_text)
            
            # æŒ‰æ®µè½åˆ†å‰²ï¼ˆä»¥åŒæ¢è¡Œæˆ–å•æ¢è¡Œ+ç¼©è¿›ä¸ºåˆ†éš”ï¼‰
            raw_paragraphs = re.split(r'\n\s*\n|\n(?=[A-Z])', full_text)
            
            for idx, para_text in enumerate(raw_paragraphs):
                text = para_text.strip()
                
                # è·³è¿‡å™ªéŸ³æ®µè½
                if self.is_noise(text):
                    continue
                
                # å¦‚æœè¿˜æ²¡æœ‰æ ‡é¢˜ï¼Œç¬¬ä¸€ä¸ªæ®µè½ä½œä¸ºæ ‡é¢˜
                if not title and text:
                    title = text[:200]  # PDFæ ‡é¢˜å¯èƒ½è¾ƒé•¿ï¼Œæˆªå–å‰200å­—ç¬¦
                    continue
                
                # åˆ¤æ–­æ˜¯å¦ä¸ºæ ‡é¢˜ï¼ˆå…¨å¤§å†™æˆ–è¾ƒçŸ­ä¸”ä»¥æ•°å­—/ç½—é©¬æ•°å­—å¼€å¤´ï¼‰
                is_heading = False
                heading_level = 0
                
                if len(text) < 100 and (text.isupper() or re.match(r'^[\dIVXLCDM]+\.\s+', text)):
                    is_heading = True
                    heading_level = 1
                
                paragraphs.append({
                    "paragraph_id": paragraph_id,
                    "text": text,
                    "type": "heading" if is_heading else "paragraph",
                    "heading_level": heading_level,
                    "original_index": idx
                })
                
                paragraph_id += 1
            
            logger.info(f"Parsed PDF: {len(paragraphs)} paragraphs, title: {title[:50]}")
            
            return {
                "title": title,
                "total_paragraphs": len(paragraphs),
                "paragraphs": paragraphs,
                "images": [],  # PDFæš‚ä¸æ”¯æŒå›¾ç‰‡æå–ï¼Œè¿”å›ç©ºæ•°ç»„
                "total_tables": 0,  # PDFè¡¨æ ¼æ•°é‡
                "format": "pdf"
            }
            
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise Exception(f"PDFè§£æå¤±è´¥: {str(e)}")
    
    def parse_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        æ ¹æ®æ–‡ä»¶ç±»å‹è§£ææ–‡æ¡£
        
        Args:
            file_content: æ–‡ä»¶äºŒè¿›åˆ¶å†…å®¹
            filename: æ–‡ä»¶å
            
        Returns:
            ç»“æ„åŒ–æ–‡æ¡£æ•°æ®
        """
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext in ['docx', 'doc']:
            return self.parse_docx(file_content)
        elif file_ext == 'pdf':
            return self.parse_pdf(file_content)
        else:
            raise ValueError(f"ä¸æ”¯æŒçš„æ–‡ä»¶æ ¼å¼: {file_ext}ï¼Œç›®å‰ä»…æ”¯æŒ .docx, .doc, .pdf")


document_parser_service = DocumentParserService()

