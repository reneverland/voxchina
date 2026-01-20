"""
文档解析模块
支持 DOCX、DOC、PDF 格式
速度优化：截断非核心部分（References、注释等）
"""

import os
import re
from pathlib import Path
from typing import Optional

# 文档解析库
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX support disabled.")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("Warning: PyPDF2 not installed. PDF support disabled.")

try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False
    print("Warning: textract not installed. DOC support may be limited.")


def extract_text_from_doc(file_path: str) -> str:
    """
    从文档中提取文本（支持 DOCX、PDF、DOC）
    
    Args:
        file_path: 文档路径
    
    Returns:
        提取的文本内容
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    ext = file_path.suffix.lower()
    
    if ext == '.docx':
        return _extract_from_docx(file_path)
    elif ext == '.pdf':
        return _extract_from_pdf(file_path)
    elif ext == '.doc':
        return _extract_from_doc_legacy(file_path)
    elif ext == '.txt':
        return _extract_from_txt(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}。支持的格式: .docx, .pdf, .doc, .txt")


def _extract_from_docx(file_path: Path) -> str:
    """从 DOCX 提取文本"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 未安装。请运行: pip install python-docx")
    
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    
    return '\n'.join(paragraphs)


def _extract_from_pdf(file_path: Path) -> str:
    """从 PDF 提取文本"""
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 未安装。请运行: pip install PyPDF2")
    
    text_parts = []
    
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        num_pages = len(pdf_reader.pages)
        
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
    
    return '\n'.join(text_parts)


def _extract_from_doc_legacy(file_path: Path) -> str:
    """从旧版 DOC 提取文本（使用 textract 或 fallback）"""
    if TEXTRACT_AVAILABLE:
        try:
            text = textract.process(str(file_path)).decode('utf-8')
            return text
        except Exception as e:
            print(f"textract 失败: {e}")
    
    # Fallback: 尝试 python-docx（部分 .doc 可能兼容）
    if DOCX_AVAILABLE:
        try:
            return _extract_from_docx(file_path)
        except Exception:
            pass
    
    raise RuntimeError(
        f"无法提取 .doc 文件内容。请安装 textract 或转换为 .docx/.pdf 格式。\n"
        f"安装命令: pip install textract"
    )


def _extract_from_txt(file_path: Path) -> str:
    """从纯文本文件提取文本"""
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def truncate_non_core_sections(
    text: str,
    mode: str = 'moderate'
) -> str:
    """
    截断非核心部分以加速处理
    
    Args:
        text: 原始文本
        mode: 截断模式
            - 'aggressive': 激进截断（单篇摘要，速度优先）
            - 'moderate': 适度截断（多篇要点提取）
            - 'none': 不截断
    
    Returns:
        截断后的文本
    """
    if mode == 'none':
        return text
    
    # 定位 References 开始位置
    ref_patterns = [
        r'\n\s*References\s*\n',
        r'\n\s*REFERENCES\s*\n',
        r'\n\s*参考文献\s*\n',
        r'\n\s*Bibliography\s*\n'
    ]
    
    ref_position = None
    for pattern in ref_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            ref_position = match.start()
            break
    
    # 截断 References 之后的内容
    if ref_position and ref_position > 1000:  # 确保有足够正文
        text = text[:ref_position]
    
    if mode == 'aggressive':
        # 激进模式：进一步限制长度
        max_chars = 30000  # 约 10k 中文字或 15k 英文单词
        if len(text) > max_chars:
            # 保留前 80% + 后 20%（保留结论部分）
            cut_point = int(max_chars * 0.8)
            tail_start = len(text) - int(max_chars * 0.2)
            text = text[:cut_point] + "\n\n[...中间内容已省略...]\n\n" + text[tail_start:]
    
    elif mode == 'moderate':
        # 适度模式：限制更宽松
        max_chars = 50000
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[...后续内容已截断...]"
    
    return text


def estimate_word_count(text: str) -> dict:
    """
    估算文本的字数（中英文分别统计）
    
    Returns:
        {
            'chinese_chars': int,
            'english_words': int,
            'total_chars': int
        }
    """
    # 中文字符（CJK Unicode 范围）
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    
    # 英文单词（简单分词）
    english_text = re.sub(r'[\u4e00-\u9fff]', ' ', text)
    english_words = len(re.findall(r'\b[a-zA-Z]+\b', english_text))
    
    return {
        'chinese_chars': chinese_chars,
        'english_words': english_words,
        'total_chars': len(text)
    }


if __name__ == "__main__":
    # 测试
    test_file = "test_data/sample.pdf"
    
    if os.path.exists(test_file):
        print(f"测试文件: {test_file}")
        
        text = extract_text_from_doc(test_file)
        print(f"提取字符数: {len(text)}")
        print(f"前200字符:\n{text[:200]}\n")
        
        text_truncated = truncate_non_core_sections(text, mode='aggressive')
        print(f"截断后字符数: {len(text_truncated)}")
        
        word_count = estimate_word_count(text)
        print(f"字数统计: {word_count}")
    else:
        print(f"测试文件不存在: {test_file}")
        print("跳过测试")
