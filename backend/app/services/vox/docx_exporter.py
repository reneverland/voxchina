"""
DOCX 导出模块
根据模式（单篇/多篇）生成格式化的 DOCX 文档
"""

import os
import re
from pathlib import Path
from typing import Optional, List, Dict

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX export disabled.")


def export_to_docx(
    content: str,
    title: str,
    mode: str,
    output_path: str = "output.docx",
    images: Optional[List[Dict[str, str]]] = None
) -> str:
    """
    导出为 DOCX 文档
    
    Args:
        content: 生成的内容文本
        title: 标题
        mode: 模式 ('SINGLE_SUMMARY' | 'MULTI_SCRIPT')
        output_path: 输出路径
        images: 提取的图片列表
    
    Returns:
        实际保存的文件路径
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx 未安装。请运行: pip install python-docx")
    
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    
    if mode == 'SINGLE_SUMMARY':
        _format_single_summary(doc, content, title, images)
    elif mode == 'MULTI_SCRIPT':
        _format_multi_script(doc, content, title, images)
    else:
        # 默认格式：纯文本
        _format_plain_text(doc, content, title)
    
    # 保存文档
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(str(output_path))
    
    return str(output_path.absolute())


def _format_single_summary(doc: Document, content: str, title: str, images: Optional[List[Dict[str, str]]] = None):
    """格式化单篇摘要"""
    
    # 标题
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 添加分隔线
    doc.add_paragraph('_' * 60)
    
    # 摘要内容（按段落分割）
    paragraphs = content.split('\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        
        para = doc.add_paragraph(para_text)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(6)
        
        # 如果是"研究发现："开头，加粗
        if para_text.startswith('研究发现：'):
            para.runs[0].font.bold = True

    # 附录图片
    _append_images(doc, images)


def _format_multi_script(doc: Document, content: str, title: str, images: Optional[List[Dict[str, str]]] = None):
    """格式化多篇口播稿（结构化）"""
    
    # 解析内容结构
    lines = content.split('\n')
    
    # 第一行：标题
    if title:
        main_title = title
    else:
        main_title = lines[0].strip() if lines else "VoxChina 口播稿"
    
    heading = doc.add_heading(main_title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()  # 空行
    
    # 逐行处理
    skip_first_title = not title  # 如果没有单独标题，跳过第一行
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 跳过第一行标题（如果已经作为 heading 添加）
        if skip_first_title:
            skip_first_title = False
            continue
        
        # 检测结构化元素
        if _is_section_heading(line):
            # 小标题（创新引领、资源协调、风险管控）
            heading2 = doc.add_heading(line, level=2)
            for run in heading2.runs:
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(51, 102, 153)
        
        elif line.startswith('Figure '):
            # Figure 占位（斜体、灰色）
            para = doc.add_paragraph(line)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.size = Pt(10)
        
        elif '欢迎继续关注' in line or 'VOXCHINA' in line or '下期再见' in line:
            # 结尾句（居中、加粗）
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(12)
        
        elif line.startswith('大家好') or '很高兴在VOXCHINA' in line:
            # 开场句（稍微突出）
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.size = Pt(11)
        
        else:
            # 普通段落
            para = doc.add_paragraph(line)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = Inches(0.25)  # 首行缩进

    # 附录图片
    _append_images(doc, images)


def _append_images(doc: Document, images: Optional[List[Dict[str, str]]]):
    """在文档末尾添加图片"""
    if not images:
        return
        
    doc.add_page_break()
    heading = doc.add_heading("附录：提取的图表/图片", level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    for img in images:
        if 'path' in img and os.path.exists(img['path']):
            try:
                # 尝试添加图片
                doc.add_picture(img['path'], width=Inches(6.0))
                # 添加文件名说明
                caption = doc.add_paragraph(f"文件: {img.get('filename', '')}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].font.size = Pt(9)
                caption.runs[0].font.italic = True
                
                doc.add_paragraph()  # 空行
            except Exception as e:
                print(f"Warning: Could not add image {img['path']}: {e}")



def _format_plain_text(doc: Document, content: str, title: str):
    """纯文本格式（fallback）"""
    if title:
        doc.add_heading(title, level=1)
    
    paragraphs = content.split('\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)


def _is_section_heading(line: str) -> bool:
    """判断是否是小标题"""
    # 匹配模式：
    # - "创新引领：..."
    # - "资源协调：..."
    # - "风险管控：..."
    # - 或其他以"："结尾的短句
    
    patterns = [
        r'^创新引领[：:]',
        r'^资源协调[：:]',
        r'^风险管控[：:]',
        r'^[一二三四五六七八九十\d]+[、\.．]\s*.{2,15}[：:]?$',  # 数字标题
    ]
    
    for pattern in patterns:
        if re.match(pattern, line):
            return True
    
    # 短句 + 冒号
    if len(line) <= 20 and ('：' in line or ':' in line):
        return True
    
    return False


if __name__ == "__main__":
    # 测试
    test_content_single = """这是一篇测试论文。作者来自某大学，研究了某个问题，使用了某种方法。
研究发现：结论1、结论2、结论3。这一结果主要归因于某种机制。"""
    
    test_content_multi = """测试口播稿标题

大家好，我是VoxChina的测试主播。
很高兴在VOXCHINA和大家见面。

今天我们从三个维度探讨：创新、协调与管控。

创新引领：
创新是第一动力。研究表明，技术创新能够推动发展。

资源协调：
资源配置至关重要。协调机制提升效率。

Figure 1. 创新与协调关系图

风险管控：
风险管控不可忽视。需要建立完善体系。

Figure 2. 风险管控框架

欢迎继续关注 VOXCHINA，我们下期再见！"""
    
    if DOCX_AVAILABLE:
        print("测试单篇摘要导出...")
        path1 = export_to_docx(
            content=test_content_single,
            title="测试论文摘要",
            mode="SINGLE_SUMMARY",
            output_path="test_output/test_single.docx"
        )
        print(f"已保存: {path1}")
        
        print("\n测试多篇口播导出...")
        path2 = export_to_docx(
            content=test_content_multi,
            title="",  # 标题在 content 中
            mode="MULTI_SCRIPT",
            output_path="test_output/test_multi.docx"
        )
        print(f"已保存: {path2}")
    else:
        print("python-docx 未安装，跳过测试")
